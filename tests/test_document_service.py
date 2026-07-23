import io
import tempfile
import unittest
from datetime import date, datetime
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import Mock, patch

from docx import Document
from openpyxl import Workbook

from document_service import DocumentService
from memory_store import MemoryStore


def make_docx_bytes():
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("青甘自驾计划", level=1)
    document.add_paragraph("8月16日从西宁前往青海湖二郎剑景区。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "住宿"
    table.cell(0, 1).text = "茶卡镇"
    document.save(buffer)
    return buffer.getvalue()


def make_xlsx_bytes():
    buffer = io.BytesIO()
    workbook = Workbook()
    itinerary = workbook.active
    itinerary.title = "每日行程"
    itinerary.append(["日期", "行程", "人数", "确认", "出发时间"])
    itinerary.append([
        date(2026, 8, 17),
        "西宁 → 青海湖 → 茶卡盐湖 → 都兰",
        4,
        True,
        datetime(2026, 8, 17, 7, 30),
    ])
    itinerary.merge_cells("A4:B4")
    itinerary["A4"] = "集合地点：西宁"

    hidden = workbook.create_sheet("隐藏备注")
    hidden.sheet_state = "hidden"
    hidden["A1"] = "不应进入知识库"
    workbook.create_sheet("空表")
    lodging = workbook.create_sheet("住宿")
    lodging.append(["日期", "住宿地"])
    lodging.append([date(2026, 8, 17), "都兰"])

    workbook.save(buffer)
    workbook.close()
    return buffer.getvalue()


class DocumentServiceTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        store = MemoryStore(Path(self.temp_dir.name) / "memory.db")
        self.service = DocumentService(store)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_extracts_docx_paragraphs_and_tables(self):
        text = self.service._extract_text("plan.docx", make_docx_bytes())

        self.assertIn("青甘自驾计划", text)
        self.assertIn("青海湖二郎剑景区", text)
        self.assertIn("住宿 | 茶卡镇", text)

    def test_extracts_visible_xlsx_rows_with_normalized_dates(self):
        text = self.service._extract_text("plan.xlsx", make_xlsx_bytes())

        self.assertIn("[工作表：每日行程]", text)
        self.assertIn("日期 | 行程 | 人数 | 确认 | 出发时间", text)
        self.assertIn("2026-08-17", text)
        self.assertIn("2026-08-17 07:30", text)
        self.assertIn("西宁 → 青海湖 → 茶卡盐湖 → 都兰", text)
        self.assertIn("4 | TRUE", text)
        self.assertEqual(text.count("集合地点：西宁"), 1)
        self.assertIn("[工作表：住宿]", text)
        self.assertIn("2026-08-17 | 都兰", text)
        self.assertNotIn("隐藏备注", text)
        self.assertNotIn("不应进入知识库", text)
        self.assertNotIn("[工作表：空表]", text)

    def test_xlsx_requests_saved_formula_results_without_recalculation(self):
        worksheet = SimpleNamespace(
            title="统计",
            sheet_state="visible",
            iter_rows=lambda values_only: iter([("总里程", 450)]),
        )
        workbook = SimpleNamespace(
            worksheets=[worksheet],
            close=Mock(),
        )
        with patch(
                "document_service.load_workbook",
                return_value=workbook) as load:
            text = self.service._extract_text("plan.xlsx", b"xlsx")

        self.assertEqual(text, "[工作表：统计]\n总里程 | 450")
        self.assertTrue(load.call_args.kwargs["read_only"])
        self.assertTrue(load.call_args.kwargs["data_only"])
        workbook.close.assert_called_once_with()

    def test_corrupt_xlsx_is_rejected_without_partial_text(self):
        with self.assertRaisesRegex(ValueError, "Excel 文件.*无法读取"):
            self.service._extract_text("broken.xlsx", b"not-a-workbook")

    def test_ingests_docx_attachment(self):
        attachment = SimpleNamespace(
            filename="plan.docx",
            url="https://example.test/plan.docx",
            size=100,
        )
        with patch.object(
                self.service,
                "_download_attachment",
                return_value=make_docx_bytes()):
            result = self.service.ingest_attachments(
                "group",
                "member",
                [attachment],
            )

        self.assertTrue(result.handled)
        self.assertIn("已保存旅行文档", result.reply)
        context = self.service.memory_store.build_document_context(
            "group",
            "茶卡住哪里",
        )
        self.assertIn("茶卡镇", context)

    def test_prepare_attachments_does_not_persist_document(self):
        attachment = SimpleNamespace(
            filename="plan.docx",
            url="https://example.test/plan.docx",
            size=100,
        )
        with patch.object(
                self.service,
                "_download_attachment",
                return_value=make_docx_bytes()):
            prepared = self.service.prepare_attachments([attachment])

        self.assertEqual(len(prepared), 1)
        self.assertEqual(prepared[0].filename, "plan.docx")
        context = self.service.memory_store.build_document_context(
            "group",
            "茶卡住哪里",
        )
        self.assertEqual(context, "")

    def test_ingests_xlsx_into_group_document_context(self):
        attachment = SimpleNamespace(
            filename="plan.xlsx",
            url="https://example.test/plan.xlsx",
            size=100,
        )
        with patch.object(
                self.service,
                "_download_attachment",
                return_value=make_xlsx_bytes()):
            result = self.service.ingest_attachments(
                "group-xlsx",
                "member",
                [attachment],
            )

        self.assertTrue(result.handled)
        self.assertIn("已保存旅行文档", result.reply)
        context = self.service.memory_store.build_document_context(
            "group-xlsx",
            "茶卡盐湖",
        )
        self.assertIn("茶卡盐湖", context)

    def test_prepare_attachments_accepts_xlsx_without_persisting(self):
        attachment = SimpleNamespace(
            filename="plan.xlsx",
            url="https://example.test/plan.xlsx",
            size=100,
        )
        with patch.object(
                self.service,
                "_download_attachment",
                return_value=make_xlsx_bytes()):
            prepared = self.service.prepare_attachments([attachment])

        self.assertEqual(len(prepared), 1)
        self.assertEqual(prepared[0].filename, "plan.xlsx")
        self.assertIn("茶卡盐湖", prepared[0].full_text)
        self.assertEqual(
            self.service.memory_store.build_document_context(
                "group",
                "茶卡盐湖",
            ),
            "",
        )

    def test_legacy_doc_requires_conversion(self):
        attachment = SimpleNamespace(
            filename="old-plan.doc",
            url="https://example.test/old-plan.doc",
            size=100,
        )

        result = self.service.ingest_attachments(
            "group",
            "member",
            [attachment],
        )

        self.assertTrue(result.handled)
        self.assertIn("转换为 .docx", result.reply)

    def test_legacy_xls_requires_conversion_to_xlsx(self):
        attachment = SimpleNamespace(
            filename="old-plan.xls",
            url="https://example.test/old-plan.xls",
            size=100,
        )

        result = self.service.ingest_attachments(
            "group",
            "member",
            [attachment],
        )

        self.assertTrue(result.handled)
        self.assertIn("另存为 .xlsx", result.reply)

    def test_document_summary_is_saved_as_long_term_context(self):
        store = MemoryStore(Path(self.temp_dir.name) / "summary.db")
        service = DocumentService(
            store,
            summarizer=lambda filename, text: "8月16日住茶卡镇，8月17日前往大柴旦。",
        )
        attachment = SimpleNamespace(
            filename="plan.docx",
            url="https://example.test/plan.docx",
            size=100,
        )
        with patch.object(
                service,
                "_download_attachment",
                return_value=make_docx_bytes()):
            result = service.ingest_attachments(
                "group-summary",
                "member",
                [attachment],
            )

        self.assertIn("已生成长期行程摘要", result.reply)
        context = store.build_document_context(
            "group-summary",
            "我们整体怎么安排？",
        )
        self.assertIn("8月17日前往大柴旦", context)


if __name__ == "__main__":
    unittest.main()
