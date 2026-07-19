from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable
from urllib.parse import urlparse

import requests
from docx import Document

from memory_store import MemoryStore


MAX_DOCUMENT_BYTES = 5 * 1024 * 1024
DOCUMENT_TIMEOUT = 20
CHUNK_SIZE = 1800
CHUNK_OVERLAP = 200
SUPPORTED_EXTENSIONS = {".docx", ".txt", ".md"}


@dataclass(frozen=True)
class DocumentIngestResult:
    handled: bool
    reply: str = ""
    memory_content: str = ""


@dataclass(frozen=True)
class PreparedDocument:
    filename: str
    sha256: str
    full_text: str
    chunks: tuple[str, ...]
    summary: str


class DocumentService:
    def __init__(
            self,
            memory_store: MemoryStore,
            summarizer: Callable[[str, str], str] | None = None):
        self.memory_store = memory_store
        self.summarizer = summarizer

    def prepare_attachments(
            self,
            attachments: list) -> tuple[PreparedDocument, ...]:
        prepared = []
        for attachment in attachments:
            filename = str(getattr(attachment, "filename", "") or "")
            if Path(filename).suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            data = self._download_attachment(attachment)
            text = self._extract_text(filename, data)
            summary = ""
            if self.summarizer:
                try:
                    summary = self.summarizer(filename, text)
                except Exception:
                    summary = ""
            prepared.append(PreparedDocument(
                filename=filename,
                sha256=hashlib.sha256(data).hexdigest(),
                full_text=text,
                chunks=tuple(self._chunk_text(text)),
                summary=summary,
            ))
        return tuple(prepared)

    def ingest_attachments(
            self,
            group_openid: str,
            member_openid: str,
            attachments: list) -> DocumentIngestResult:
        document_attachments = []
        legacy_documents = []

        for attachment in attachments:
            filename = str(getattr(attachment, "filename", "") or "")
            extension = Path(filename).suffix.lower()
            if extension in SUPPORTED_EXTENSIONS:
                document_attachments.append(attachment)
            elif extension == ".doc":
                legacy_documents.append(filename)

        if not document_attachments and not legacy_documents:
            return DocumentIngestResult(handled=False)

        messages = []
        memory_names = []

        for filename in legacy_documents:
            memory_names.append(filename)
            messages.append(
                f"暂不支持旧版 Word 文件 {filename}，请转换为 .docx 后重新上传。"
            )

        for attachment in document_attachments:
            filename = str(attachment.filename)
            try:
                data = self._download_attachment(attachment)
                text = self._extract_text(filename, data)
                chunks = self._chunk_text(text)
                digest = hashlib.sha256(data).hexdigest()
                summary = ""
                summary_failed = False
                if self.summarizer:
                    try:
                        summary = self.summarizer(filename, text)
                    except Exception:
                        summary_failed = True
                stored = self.memory_store.add_document(
                    group_openid=group_openid,
                    uploader_openid=member_openid,
                    filename=filename,
                    sha256=digest,
                    full_text=text,
                    chunks=chunks,
                    summary=summary,
                )
                memory_names.append(filename)
                if stored.is_new:
                    messages.append(
                        f"已保存旅行文档：{filename}（{len(text)} 字，{len(chunks)} 个片段）。"
                    )
                    if summary:
                        messages.append("已生成长期行程摘要。")
                    elif summary_failed:
                        messages.append("自动摘要失败，但文档原文和分块已正常保存。")
                else:
                    messages.append(f"该文档已保存过：{stored.filename}。")
            except Exception as exc:
                messages.append(f"处理文档 {filename} 失败：{exc}")

        messages.append(
            "文档属于群共享长期资料，不受最近 6 轮对话限制。后续提问时会按内容检索相关片段。"
        )
        return DocumentIngestResult(
            handled=True,
            reply="\n".join(messages),
            memory_content="上传旅行文档：" + "、".join(memory_names),
        )

    @staticmethod
    def _download_attachment(attachment) -> bytes:
        url = str(getattr(attachment, "url", "") or "")
        parsed = urlparse(url)
        if parsed.scheme != "https" or not parsed.netloc:
            raise ValueError("附件下载地址不是有效的 HTTPS URL")

        declared_size = int(getattr(attachment, "size", 0) or 0)
        if declared_size > MAX_DOCUMENT_BYTES:
            raise ValueError("文件超过 5 MB 限制")

        response = requests.get(url, timeout=DOCUMENT_TIMEOUT)
        response.raise_for_status()
        data = response.content
        if len(data) > MAX_DOCUMENT_BYTES:
            raise ValueError("文件超过 5 MB 限制")
        return data

    @staticmethod
    def _extract_text(filename: str, data: bytes) -> str:
        extension = Path(filename).suffix.lower()
        if extension == ".docx":
            document = Document(io.BytesIO(data))
            parts = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
        else:
            text = DocumentService._decode_text(data)

        text = text.replace("\x00", "")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if len(text) < 10:
            raise ValueError("文档中没有提取到足够的文本")
        return text

    @staticmethod
    def _decode_text(data: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("文本文件编码无法识别")

    @staticmethod
    def _chunk_text(text: str) -> list[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = min(len(text), start + CHUNK_SIZE)
            if end < len(text):
                boundary = max(
                    text.rfind("\n", start, end),
                    text.rfind("。", start, end),
                )
                if boundary > start + CHUNK_SIZE // 2:
                    end = boundary + 1

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            if end >= len(text):
                break
            start = max(start + 1, end - CHUNK_OVERLAP)

        return chunks
