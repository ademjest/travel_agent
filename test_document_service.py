import io
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document

from document_service import DocumentService
from memory_store import MemoryStore


def make_docx_bytes():
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("青甘自驾计划", level=1)
    document.add_paragraph("8月16日从西宁前往青海湖二郎剑景区。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "住宿"
    table.cell(0, 1).text = "茶卡镇"
    document.save(buffer)
    return buffer.getvalue()


class DocumentServiceTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        store = MemoryStore(Path(self.temp_dir.name) / "memory.db")
        self.service = DocumentService(store)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_extracts_docx_paragraphs_and_tables(self):
        text = self.service._extract_text("plan.docx", make_docx_bytes())

        self.assertIn("青甘自驾计划", text)
        self.assertIn("青海湖二郎剑景区", text)
        self.assertIn("住宿 | 茶卡镇", text)

    def test_ingests_docx_attachment(self):
        attachment = SimpleNamespace(
            filename="plan.docx",
            url="https://example.test/plan.docx",
            size=100,
        )
        with patch.object(
                self.service,
                "_download_attachment",
                return_value=make_docx_bytes()):
            result = self.service.ingest_attachments(
                "group",
                "member",
                [attachment],
            )

        self.assertTrue(result.handled)
        self.assertIn("已保存旅行文档", result.reply)
        context = self.service.memory_store.build_document_context(
            "group",
            "茶卡住哪里",
        )
        self.assertIn("茶卡镇", context)

    def test_legacy_doc_requires_conversion(self):
        attachment = SimpleNamespace(
            filename="old-plan.doc",
            url="https://example.test/old-plan.doc",
            size=100,
        )

        result = self.service.ingest_attachments(
            "group",
            "member",
            [attachment],
        )

        self.assertTrue(result.handled)
        self.assertIn("转换为 .docx", result.reply)

    def test_document_summary_is_saved_as_long_term_context(self):
        store = MemoryStore(Path(self.temp_dir.name) / "summary.db")
        service = DocumentService(
            store,
            summarizer=lambda filename, text: "8月16日住茶卡镇，8月17日前往大柴旦。",
        )
        attachment = SimpleNamespace(
            filename="plan.docx",
            url="https://example.test/plan.docx",
            size=100,
        )
        with patch.object(
                service,
                "_download_attachment",
                return_value=make_docx_bytes()):
            result = service.ingest_attachments(
                "group-summary",
                "member",
                [attachment],
            )

        self.assertIn("已生成长期行程摘要", result.reply)
        context = store.build_document_context(
            "group-summary",
            "我们整体怎么安排？",
        )
        self.assertIn("8月17日前往大柴旦", context)


if __name__ == "__main__":
    unittest.main()
